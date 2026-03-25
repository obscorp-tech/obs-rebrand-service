"""Tests for FastAPI service."""

from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

from rebrand_service.api import app


@pytest.fixture
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture
def sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("Test body.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestHealthEndpoint:
    def test_health(self, client: TestClient) -> None:
        response = client.get("/health")
        assert response.status_code == 200
        data = response.json()
        assert data["status"] in ("healthy", "degraded")

    def test_clients_list(self, client: TestClient) -> None:
        response = client.get("/clients")
        assert response.status_code == 200
        assert "clients" in response.json()


class TestRebrandEndpoint:
    def test_rebrand_unknown_client(self, client: TestClient, sample_docx_bytes: bytes) -> None:
        response = client.post(
            "/rebrand/nonexistent-client",
            files={"file": ("test.docx", sample_docx_bytes)},
        )
        assert response.status_code == 404

    def test_rebrand_unsupported_filetype(self, client: TestClient) -> None:
        response = client.post(
            "/rebrand/acme-corp",
            files={"file": ("test.txt", b"not a doc")},
        )
        # Either 400 (unsupported) or 404 (client not found in test env)
        assert response.status_code in (400, 404)
