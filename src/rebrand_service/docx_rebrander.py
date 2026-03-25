"""DOCX rebranding engine — applies brand config to Word documents."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from rebrand_service.models import BrandConfig, compute_file_hash

logger = logging.getLogger(__name__)


class DocxRebrander:
    """Applies brand configuration to DOCX files idempotently."""

    def __init__(self, brand: BrandConfig, repo_root: Path | None = None) -> None:
        self.brand = brand
        self.repo_root = repo_root or Path.cwd()

    def rebrand(self, input_path: Path, output_path: Path) -> dict[str, str]:
        """
        Rebrand a DOCX file and write to output_path.

        Returns audit metadata dict with input/output hashes.
        """
        try:
            input_hash = compute_file_hash(input_path)
            logger.info(
                "Rebranding DOCX: %s -> %s (client=%s, input_sha256=%s)",
                input_path.name,
                output_path.name,
                self.brand.client_slug,
                input_hash[:16],
            )
            doc = Document(str(input_path))
            self._apply_typography(doc)
            self._apply_colors(doc)
            self._apply_logo(doc)
            self._apply_compliance_footer(doc)

            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            output_hash = compute_file_hash(output_path)
            logger.info(
                "Rebranded DOCX saved: %s (output_sha256=%s)",
                output_path.name,
                output_hash[:16],
            )

            return {
                "input_file": str(input_path),
                "output_file": str(output_path),
                "input_sha256": input_hash,
                "output_sha256": output_hash,
                "client": self.brand.client_slug,
                "status": "success",
            }

        except Exception as e:
            logger.exception("Failed to rebrand %s: %s", input_path.name, e)
            return {
                "input_file": str(input_path),
                "output_file": str(output_path),
                "client": self.brand.client_slug,
                "status": "error",
                "error": str(e),
            }

    def _apply_typography(self, doc: Document) -> None:
        """Apply font family and sizes to all runs."""
        typo = self.brand.typography

        for paragraph in doc.paragraphs:
            is_heading = paragraph.style.name.startswith("Heading")
            for run in paragraph.runs:
                run.font.name = typo.heading_font if is_heading else typo.body_font
                run.font.size = Pt(typo.heading_size_pt) if is_heading else Pt(typo.body_size_pt)
                # Force font for East Asian text
                rpr = run._element.get_or_add_rPr()
                rpr_fonts = rpr.find(qn("w:rFonts"))
                if rpr_fonts is None:
                    rpr_fonts = rpr.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rpr_fonts)
                font_name = typo.heading_font if is_heading else typo.body_font
                rpr_fonts.set(qn("w:ascii"), font_name)
                rpr_fonts.set(qn("w:hAnsi"), font_name)
                rpr_fonts.set(qn("w:cs"), font_name)

        # Also apply to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = typo.body_font
                            run.font.size = Pt(typo.body_size_pt)

    def _apply_colors(self, doc: Document) -> None:
        """Apply brand colors to text runs."""
        colors = self.brand.colors
        heading_rgb = RGBColor.from_string(colors.heading_text)
        body_rgb = RGBColor.from_string(colors.body_text)

        for paragraph in doc.paragraphs:
            is_heading = paragraph.style.name.startswith("Heading")
            for run in paragraph.runs:
                run.font.color.rgb = heading_rgb if is_heading else body_rgb

    def _apply_logo(self, doc: Document) -> None:
        """Insert or replace logo in header."""
        if self.brand.logo is None:
            return

        logo_path = self.repo_root / self.brand.logo.path
        if not logo_path.exists():
            logger.warning("Logo file not found: %s", logo_path)
            return

        if self.brand.logo.position == "header":
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False
                # Clear existing header content
                for p in header.paragraphs:
                    p.clear()
                # Add logo to first paragraph
                if header.paragraphs:
                    run = header.paragraphs[0].add_run()
                    run.add_picture(
                        str(logo_path),
                        width=Inches(self.brand.logo.width_inches),
                    )
                    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _apply_compliance_footer(self, doc: Document) -> None:
        """Add compliance footer text if configured."""
        compliance = self.brand.compliance
        footer_parts: list[str] = []

        if compliance.confidentiality_label:
            footer_parts.append(compliance.confidentiality_label)
        if compliance.footer_text:
            footer_parts.append(compliance.footer_text)

        if not footer_parts:
            return

        footer_text = " | ".join(footer_parts)

        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            # Clear and set footer
            for p in footer.paragraphs:
                p.clear()
            if footer.paragraphs:
                run = footer.paragraphs[0].add_run(footer_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string("999999")
                footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
